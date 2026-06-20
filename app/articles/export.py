import os

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.articles.ieee_template import clear_document_body, open_template, set_columns


def _style(doc, name, fallback="Normal"):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles[fallback]


def _add_authors(doc, authors, affiliations):
    author_style = _style(doc, "Author")
    p = doc.add_paragraph(style=author_style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, author in enumerate(authors):
        if i > 0:
            p.add_run(", ")
        p.add_run(author["name"])
        sup = p.add_run(str(author["aff"]))
        sup.font.superscript = True

    for aff in affiliations:
        aff_p = doc.add_paragraph(style=author_style)
        aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        marker = aff_p.add_run("^1 ")
        marker.font.superscript = True
        aff_p.add_run(aff.replace("^1 ", ""))


def _add_abstract(doc, article):
    p = doc.add_paragraph(style=_style(doc, "Abstract"))
    label = p.add_run("Abstract—")
    label.bold = True
    body = p.add_run(article["abstract"])
    body.italic = True

    terms = doc.add_paragraph(style=_style(doc, "Keywords"))
    terms_label = terms.add_run("Index Terms—")
    terms_label.bold = True
    terms.add_run(article["index_terms"])


def _add_section_heading(doc, sec):
    doc.add_paragraph(f"{sec['id']}. {sec['title']}", style=_style(doc, "Heading 1"))


def _add_body_text(doc, text):
    body_style = _style(doc, "Body Text")
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        p = doc.add_paragraph(style=body_style)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(block)


def _add_figure(doc, figure, upload_folder, width_inches=3.25):
    path = os.path.join(upload_folder, figure["file_path"])
    if not os.path.exists(path):
        return

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(6)
    img_p.paragraph_format.space_after = Pt(2)
    run = img_p.add_run()
    run.add_picture(path, width=Inches(width_inches))

    cap = doc.add_paragraph(style=_style(doc, "figure caption"))
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(f"Fig. {figure['num']}. {figure['caption']}")


def _add_references(doc, references):
    doc.add_paragraph("References", style=_style(doc, "Heading 5"))
    ref_style = _style(doc, "references")
    for ref in references:
        p = doc.add_paragraph(style=ref_style)
        p.add_run(f"[{ref['num']}] {ref['text']}")


def build_article_document(article, upload_folder, template_path):
    """Preenche o template IEEE Word e retorna o Document."""
    doc, _tmp = open_template(template_path)
    clear_document_body(doc)
    set_columns(doc.sections[-1], num=1)

    title_p = doc.add_paragraph(article["title"], style=_style(doc, "paper title"))
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_authors(doc, article["authors"], article["affiliations"])
    _add_abstract(doc, article)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    body_section.top_margin = doc.sections[0].top_margin
    set_columns(body_section, num=2, space_twips=360)

    figures = article.get("figures") or []
    for sec in article["sections"]:
        _add_section_heading(doc, sec)
        _add_body_text(doc, sec["content"])
        if sec["id"] == "IV" and figures:
            for figure in figures:
                _add_figure(doc, figure, upload_folder)

    participants = article.get("participants") or []
    if participants:
        doc.add_paragraph("Authors and Participants", style=_style(doc, "Heading 5"))
        body_style = _style(doc, "Body Text")
        for item in participants:
            student = item["student"]
            line = (
                f"{student.name} — {item['role']}, {item['group']}, "
                f"turma {student.class_name} ({student.email})"
            )
            p = doc.add_paragraph(style=body_style)
            p.add_run("• " + line)

    doc.add_paragraph("Acknowledgment", style=_style(doc, "Heading 5"))
    ack = doc.add_paragraph(style=_style(doc, "Body Text"))
    ack.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ack_run = ack.add_run(article["acknowledgment"])
    ack_run.italic = True

    _add_references(doc, article["references"])
    return doc


def export_article_docx(article, project, upload_folder, template_path=None):
    if template_path is None:
        from flask import current_app

        template_path = current_app.config.get("IEEE_ARTICLE_TEMPLATE")

    if template_path and os.path.isfile(template_path):
        from io import BytesIO

        doc = build_article_document(article, upload_folder, template_path)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    raise FileNotFoundError(
        "Template IEEE não encontrado. Coloque conference-template-a4.docx em arquivos/."
    )
